import json
import os
import re
import time
import difflib
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.shared import Inches, RGBColor
from openai import OpenAI
import streamlit as st
from pinecone import Pinecone, ServerlessSpec

# =====================================================================
# 1. APPLICATION SETUP & PINECONE DB CONFIG
# =====================================================================
st.set_page_config(
    page_title="NVIDIA AI Legal Reviewer (Borrower-Friendly Review)",
    page_icon="💬",
    layout="wide",
)

INDEX_NAME = "project-finance-playbook"
EMBED_MODEL = "multilingual-e5-large"
NVIDIA_BASE_URL = "https://integrate.api.nvidia.com/v1"

# Active Model Endpoint
NVIDIA_MODEL = "meta/llama-3.1-8b-instruct"

# Batch size for LLM inference calls
BATCH_SIZE = 4 

@st.cache_resource
def init_pinecone(api_key):
    """Initializes Pinecone client and ensures serverless index exists."""
    if not api_key:
        return None, None
    try:
        pc = Pinecone(api_key=api_key)
        existing_indexes = [idx.name for idx in pc.list_indexes()]
        if INDEX_NAME not in existing_indexes:
            pc.create_index(
                name=INDEX_NAME,
                dimension=1024,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1")
            )
        index = pc.Index(INDEX_NAME)
        return pc, index
    except Exception as e:
        st.error(f"Failed to initialize Pinecone: {str(e)}")
        return None, None

# =====================================================================
# 2. WORD DOCUMENT PARSING & REDLINED SIDEBAR COMMENT GENERATOR
# =====================================================================

def extract_paragraphs_from_docx(docx_file):
    doc = Document(docx_file)
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text and len(text.split()) >= 4:
            paragraphs.append(text)
    return paragraphs


def add_redlined_comment_content(comment_paragraph, orig_text, suggested_text, explanation):
    """
    Constructs a word-by-word redlined comparison inside the balloon comment paragraph.
    - Deleted text: Cut in Red with Strike-through
    - Inserted text: Green text
    - Unchanged text: Standard font
    """
    r_hdr1 = comment_paragraph.add_run("💡 BORROWER-FRIENDLY REDLINE:\n")
    r_hdr1.bold = True
    
    orig_words = orig_text.split()
    sugg_words = suggested_text.split()
    matcher = difflib.SequenceMatcher(None, orig_words, sugg_words)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            comment_paragraph.add_run(" ".join(orig_words[i1:i2]) + " ")
        elif tag == 'delete':
            del_run = comment_paragraph.add_run(" ".join(orig_words[i1:i2]) + " ")
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(218, 41, 28)  # Red
        elif tag == 'insert':
            ins_run = comment_paragraph.add_run(" ".join(sugg_words[j1:j2]) + " ")
            ins_run.font.color.rgb = RGBColor(34, 139, 34)  # Green
            ins_run.bold = True
        elif tag == 'replace':
            del_run = comment_paragraph.add_run(" ".join(orig_words[i1:i2]) + " ")
            del_run.font.strike = True
            del_run.font.color.rgb = RGBColor(218, 41, 28)  # Red
            
            ins_run = comment_paragraph.add_run(" ".join(sugg_words[j1:j2]) + " ")
            ins_run.font.color.rgb = RGBColor(34, 139, 34)  # Green
            ins_run.bold = True

    comment_paragraph.add_run("\n\n")
    r_hdr2 = comment_paragraph.add_run("📌 BORROWER PROTECTION RATIONALE:\n")
    r_hdr2.bold = True
    
    r_exp = comment_paragraph.add_run(explanation)
    r_exp.font.italic = True


def create_commented_docx(paragraph_results, author="Borrower AI Legal Counsel"):
    """
    Generates a clean DOCX where suggestions and redlined revisions are placed
    inside native MS Word sidebar balloon comment bubbles.
    """
    doc = Document()

    for orig_text, suggested_text, explanation, is_acceptable in paragraph_results:
        p = doc.add_paragraph()
        p.add_run(orig_text)

        if not is_acceptable and orig_text.strip() != suggested_text.strip():
            comment_added = False
            try:
                comment = doc.add_comment(
                    runs=p.runs,
                    author=author,
                    initials="BORROWER"
                )
                comment_p = comment.paragraphs[0]
                add_redlined_comment_content(comment_p, orig_text, suggested_text, explanation)
                comment_added = True
            except Exception as e:
                print(f"[DEBUG] Balloon comment creation notice: {e}", flush=True)

            if not comment_added:
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.left_indent = Inches(0.3)
                add_redlined_comment_content(p_sub, orig_text, suggested_text, explanation)

    output_path = "Borrower_Friendly_Facility_Agreement_Comments.docx"
    doc.save(output_path)
    return output_path

# =====================================================================
# 3. CLOUD VECTOR RETRIEVAL & BORROWER-ADVOCATE LLM ENGINE
# =====================================================================

def query_pinecone_batch(pc, index, chunk_paras, chunk_start_idx):
    """Retrieves top 5 contextual precedents per clause to check baseline terms."""
    try:
        embeddings = pc.inference.embed(
            model=EMBED_MODEL,
            inputs=chunk_paras,
            parameters={"input_type": "query"}
        )
        
        batch_results = []
        for idx, (p_text, emb) in enumerate(zip(chunk_paras, embeddings)):
            res = index.query(
                vector=emb["values"],
                top_k=5,
                include_metadata=True
            )
            
            ctx_list = []
            if res.get("matches") and len(res["matches"]) > 0:
                for match in res["matches"]:
                    if match["score"] > 0.45:
                        doc_str = match["metadata"].get("text", "")
                        src = match["metadata"].get("source", "Repo")
                        ctx_list.append(f"Precedent Chunk [{src}]: \"{doc_str}\"")
            
            ctx = "\n".join(ctx_list) if ctx_list else "NO DIRECT PINECONE PRECEDENT FOUND FOR THIS CLAUSE."
            
            batch_results.append({
                "id": chunk_start_idx + idx,
                "clause": p_text,
                "context": ctx
            })
        return batch_results, len(chunk_paras)
    except Exception as e:
        print(f"[DEBUG] ❌ Pinecone Query Error: {e}", flush=True)
        return [{
            "id": chunk_start_idx + idx,
            "clause": p_text,
            "context": "NO DIRECT PINECONE PRECEDENT FOUND FOR THIS CLAUSE."
        } for idx, p_text in enumerate(chunk_paras)], len(chunk_paras)


def run_parallel_pinecone_retrieval(pc, index, paragraphs, batch_size=5, max_workers=3, status_placeholder=None, progress_bar=None, log_area=None, logs_list=None):
    total = len(paragraphs)
    prepared_items = [None] * total
    
    chunks = [
        (paragraphs[i : i + batch_size], i) 
        for i in range(0, total, batch_size)
    ]
    
    completed_clauses = 0
    total_chunks = len(chunks)
    completed_chunks = 0

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(query_pinecone_batch, pc, index, chunk, start_idx): (chunk, start_idx)
            for chunk, start_idx in chunks
        }
        
        for future in as_completed(futures):
            batch_data, chunk_len = future.result()
            completed_clauses += chunk_len
            completed_chunks += 1
            
            for item in batch_data:
                prepared_items[item["id"]] = item

            if status_placeholder and progress_bar:
                pct = int((completed_clauses / total) * 20)
                progress_bar.progress(pct)
                status_placeholder.text(f"🔍 Searching Precedents: {completed_clauses}/{total} clauses...")
                
            if log_area and logs_list is not None:
                log_msg = f"[{time.strftime('%H:%M:%S')}] ☁️ Batch {completed_chunks}/{total_chunks} retrieved ({completed_clauses}/{total} clauses)"
                logs_list.append(log_msg)
                log_area.text("\n".join(logs_list[-12:]))

    return prepared_items


def extract_json_from_text(raw_text):
    """Extracts JSON payload safely whether the LLM outputs an Object or a List."""
    raw_text = raw_text.strip()
    
    if raw_text.startswith("```"):
        raw_text = re.sub(r"^```[a-zA-Z]*\n?", "", raw_text)
        raw_text = re.sub(r"\n?```$", "", raw_text)
        raw_text = raw_text.strip()

    match = re.search(r"(\{[\s\S]*\}|\[[\s\S]*\])", raw_text)
    if match:
        json_candidate = match.group(0)
        try:
            return json.loads(json_candidate)
        except json.JSONDecodeError:
            pass

    return json.loads(raw_text)


def analyze_clause_batch_llm(batch_items, custom_instruction, nvidia_api_key):
    """Analyzes clause batches acting specifically as Borrower's Legal Counsel to maximize operational flexibility and minimize risk."""
    if not nvidia_api_key:
        st.error("❌ API Key is missing!")
        return []

    client = OpenAI(
        base_url=NVIDIA_BASE_URL,
        api_key=nvidia_api_key,
        timeout=120.0,
        max_retries=3
    )

    system_prompt = """
    You are Senior Legal Counsel representing the BORROWER in loan agreement negotiations.

    YOUR OBJECTIVE: Protect the Borrower from onerous lender terms, remove strict/unreasonable restrictions, ensure broad operational headroom, and introduce standard market borrower protections.

    BORROWER-FRIENDLY MANDATORY REVIEW RULES:

    1. SOURCE A (PINECONE PRECEDENT CONTEXT):
       - Use retrieved precedent data to prevent the Lender from imposing financial ratios or fees tighter than baseline agreements (e.g., maintain DSCR caps, leverage limits, interest margins).
       - If draft terms are harsher on the Borrower than Source A, REJECT ('is_acceptable': false) and align with the most favorable precedent term.

    2. SOURCE B (BORROWER LEGAL DEFENSE & RISK REDUCTION):
       - MANDATORILY REJECT ('is_acceptable': false) AND REDLINE any clause that contains:
         * UNILATERAL LENDER DISCRETION: Change "in the sole discretion/opinion of the Lender" to "acting reasonably in consultation with the Borrower" or "certified by the Lender's Technical Advisor / Banking Base Case".
         * ABSENCE OF MATERIALITY QUALIFIERS: Insert "in all material respects" or "Material Adverse Effect" triggers before representation/covenant breaches.
         * MISSING CURE PERIODS: Add mandatory cure/grace windows (e.g., "30 days after written notice" for non-payment or general covenants, "5 Business Days" for financial payments).
         * UNFAIR PAYMENT SHIFTS: Change non-business day payment shifts from "immediately preceding Business Day" to "next succeeding Business Day".
         * NARROW DEFINITIONS: Expand key borrower-favorable definitions (e.g., ensure "Permitted Indebtedness" includes trade credits, working capital lines, and subordinated sponsor loans; ensure "Contracts" covers all major project agreements).
         * ABSOLUTE RESTRICTIONS: Change absolute prohibitions on asset transfers, corporate restructuring, or capex into exceptions with "prior written consent of the Lender (such consent not to be unreasonably withheld, conditioned, or delayed)".

    OUTPUT DIRECTIVES:
    - For every rejected clause, propose a complete, borrower-protective redlined text in 'proposed_text'.
    - In 'explanation', detail the borrower risk using: "[Source A Precedent Protection]" or "[Source B Borrower Protection]" followed by the rationale.
    - Respond STRICTLY in valid JSON using this structure:
    {
      "results": [
        {
          "id": 1,
          "is_acceptable": false,
          "proposed_text": "string",
          "explanation": "string"
        }
      ]
    }
    """
    
    formatted_input = [
        {
            "id": item["id"], 
            "clause": item["clause"], 
            "pinecone_repository_context_source_a": item["context"]
        }
        for item in batch_items
    ]

    user_prompt = f"BORROWER DEAL DIRECTIVES: {custom_instruction}\nCLAUSES TO REVIEW: {json.dumps(formatted_input)}"

    for attempt in range(3):
        try:
            print(f"[DEBUG] Sending LLM Request (Attempt {attempt+1})...", flush=True)
            response = client.chat.completions.create(
                model=NVIDIA_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.0
            )
            
            raw_content = response.choices[0].message.content
            data = extract_json_from_text(raw_content)
            
            # Safe handling for both Array and Dict returns
            if isinstance(data, list):
                return data
            elif isinstance(data, dict):
                return data.get("results", [])
                
        except Exception as e:
            err_msg = f"❌ [Attempt {attempt+1} Failed] Model: '{NVIDIA_MODEL}' | Error: {str(e)}"
            print(f"[DEBUG] {err_msg}", flush=True)
            time.sleep((attempt + 1) * 3)

    return [
        {
            "id": item["id"],
            "is_acceptable": True,
            "proposed_text": item["clause"],
            "explanation": "Accepted by default due to connection error."
        }
        for item in batch_items
    ]

# =====================================================================
# 4. STREAMLIT UI & TABBED INTERFACE
# =====================================================================

st.title("💬 Contract AI Auditor T-Bajaj (Borrower Advocate Review)")
st.caption("Substantive Borrower-Friendly Contract Review powered by Pinecone Precedents & NVIDIA Llama 3.1 8B")

default_nvidia = st.secrets.get("NVIDIA_API_KEY", "") if "NVIDIA_API_KEY" in st.secrets else ""
default_pinecone = st.secrets.get("PINECONE_API_KEY", "") if "PINECONE_API_KEY" in st.secrets else ""

with st.sidebar:
    st.header("🔑 Credentials")
    nvidia_api_key = st.text_input("NVIDIA API Key", value=default_nvidia, type="password")
    pinecone_api_key = st.text_input("Pinecone API Key", value=default_pinecone, type="password")
    
    st.divider()
    st.header("⚙️ Model Settings")
    st.info(f"**Active Model:** `{NVIDIA_MODEL}`\n\n**Batch Size:** `{BATCH_SIZE} Clauses / Call`")
    
    st.divider()
    st.header("☁️ Cloud Vector Storage")
    
    pc_client, pc_index = init_pinecone(pinecone_api_key)
    if pc_index:
        try:
            stats = pc_index.describe_index_stats()
            vector_count = stats.get("total_vector_count", 0)
            st.metric("Cloud Index Vector Count", vector_count)
        except Exception:
            st.caption("Unable to fetch stats.")
            
        if st.button("🗑️ Clear Cloud Memory"):
            try:
                pc_index.delete(delete_all=True)
                st.success("Cloud database erased!")
                st.rerun()
            except Exception as e:
                st.error(f"Error clearing cloud database: {e}")

tab_review, tab_repository, tab_chat = st.tabs([
    "💬 Sidebar Comment Reviewer", 
    "📚 Upload Precedent Agreements", 
    "🤖 Contract Chat Assistant"
])

# ---------------------------------------------------------------------
# TAB 1: WORD COMMENT REVIEWER (BORROWER OPTIMIZED)
# ---------------------------------------------------------------------
with tab_review:
    st.header("Review Draft Facility Agreement (Borrower Perspective) - TBajaj")
    
    uploaded_draft = st.file_uploader("Upload Target Facility Agreement (.docx)", type=["docx"], key="target_doc")
    custom_instruction = st.text_area(
        "Optional Borrower Deal Directives / Overrides", 
        value="Maximize borrower flexibility. Introduce 30-day cure periods for covenant defaults, add materiality thresholds, change 'sole discretion of Lender' to 'acting reasonably', shift non-business day payments to the next business day, and carve out working capital / trade debt under Permitted Indebtedness.",
        placeholder="e.g., 'Ensure minimum DSCR is negotiable and cure periods are added to all default clauses.'"
    )
    
    if st.button("💬 Analyze Contract (Generate Borrower-Friendly Redlines)", type="primary"):
        if not nvidia_api_key:
            st.error("Please enter your NVIDIA API Key.")
        elif not pinecone_api_key:
            st.error("Please enter your Pinecone API Key.")
        elif not uploaded_draft:
            st.error("Please upload a .docx document.")
        else:
            paragraphs = extract_paragraphs_from_docx(uploaded_draft)
            total_paragraphs = len(paragraphs)
            st.info(f"Loaded document: {total_paragraphs} clauses identified.")
            
            start_time = time.time()
            progress_bar = st.progress(0)
            status_placeholder = st.empty()
            
            debug_box = st.expander("🔍 Live Debug & Processing Logs", expanded=True)
            log_area = debug_box.empty()
            logs = []

            # STEP 1: Vector Search
            logs.append(f"[{time.strftime('%H:%M:%S')}] [DEBUG] Fetching Precedents from Pinecone...")
            log_area.text("\n".join(logs))
            
            vec_start = time.time()
            prepared_items = run_parallel_pinecone_retrieval(
                pc=pc_client,
                index=pc_index, 
                paragraphs=paragraphs, 
                batch_size=5, 
                max_workers=3,
                status_placeholder=status_placeholder,
                progress_bar=progress_bar,
                log_area=log_area,
                logs_list=logs
            )
            vec_elapsed = round(time.time() - vec_start, 2)
            
            progress_bar.progress(20)
            logs.append(f"[{time.strftime('%H:%M:%S')}] [DEBUG] Precedents retrieved in {vec_elapsed}s! Evaluating via Borrower-Advocate Engine...")
            log_area.text("\n".join(logs[-12:]))

            # STEP 2: LLM Borrower-Friendly Evaluation
            comment_results = []
            total_items = len(prepared_items)
            
            for i in range(0, total_items, BATCH_SIZE):
                batch = prepared_items[i : i + BATCH_SIZE]
                current_c = min(i + BATCH_SIZE, total_items)
                
                batch_evals = analyze_clause_batch_llm(
                    batch_items=batch, 
                    custom_instruction=custom_instruction, 
                    nvidia_api_key=nvidia_api_key
                )
                
                eval_map = {res["id"]: res for res in batch_evals if isinstance(res, dict) and "id" in res}
                for item in batch:
                    res = eval_map.get(item["id"], {
                        "is_acceptable": True,
                        "proposed_text": item["clause"],
                        "explanation": ""
                    })
                    
                    comment_results.append((
                        item["clause"],
                        res.get("proposed_text", item["clause"]),
                        res.get("explanation", ""),
                        res.get("is_acceptable", True)
                    ))
                
                pct = 20 + int((current_c / total_items) * 80)
                progress_bar.progress(pct)
                status_placeholder.text(f"Evaluated {current_c}/{total_items} clauses ({pct}%)...")
                
                log_msg = f"[{time.strftime('%H:%M:%S')}] [DEBUG] Evaluated Clauses {i+1} to {current_c} of {total_items} ({pct}%)"
                logs.append(log_msg)
                log_area.text("\n".join(logs[-12:]))
            
            elapsed = round(time.time() - start_time, 2)
            st.success(f"⚡ Completed borrower review in {elapsed} seconds!")
            
            output_docx = create_commented_docx(comment_results)
            with open(output_docx, "rb") as file_data:
                st.download_button(
                    label="📥 Download File with Borrower Redlines & Comments (.docx)",
                    data=file_data,
                    file_name="Borrower_Friendly_Facility_Agreement_Comments.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ---------------------------------------------------------------------
# TAB 2: REPOSITORY INDEXING
# ---------------------------------------------------------------------
with tab_repository:
    st.header("Upload Precedent Facility Agreements to Cloud")
    precedent_files = st.file_uploader("Upload Loan Agreements (.docx)", type=["docx"], accept_multiple_files=True)
    
    if precedent_files and st.button("☁️ Save to Pinecone Cloud"):
        if not pinecone_api_key or not pc_client:
            st.error("Please provide a valid Pinecone API key.")
        else:
            prog_bar = st.progress(0)
            status_box = st.empty()
            debug_container = st.expander("🛠️ Real-Time Cloud Indexing Logs", expanded=True)
            log_console = debug_container.empty()
            logs = []
            
            total_files = len(precedent_files)
            total_indexed = 0
            
            for file_idx, file in enumerate(precedent_files):
                paragraphs = extract_paragraphs_from_docx(file)
                total_paras = len(paragraphs)
                if total_paras == 0:
                    continue
                
                b_size = 10
                clean_fname = re.sub(r'[^a-zA-Z0-9_]', '_', file.name)
                
                for i in range(0, total_paras, b_size):
                    chunk = paragraphs[i : i + b_size]
                    
                    try:
                        embeddings = pc_client.inference.embed(
                            model=EMBED_MODEL,
                            inputs=chunk,
                            parameters={"input_type": "passage"}
                        )
                        
                        vectors_to_upsert = []
                        for j, (text, emb) in enumerate(zip(chunk, embeddings)):
                            v_id = f"{clean_fname}_p{i+j}_{int(time.time()*1000)}"
                            metadata_text = text[:1500] if len(text) > 1500 else text
                            
                            vectors_to_upsert.append({
                                "id": v_id,
                                "values": emb["values"],
                                "metadata": {
                                    "text": metadata_text, 
                                    "source": file.name, 
                                    "chunk": i+j
                                }
                            })
                        
                        pc_index.upsert(vectors=vectors_to_upsert)
                        total_indexed += len(chunk)
                        
                        current_clause = min(i + b_size, total_paras)
                        overall_pct = ((file_idx + (current_clause / total_paras)) / total_files)
                        prog_bar.progress(min(overall_pct, 1.0))
                        status_box.text(f"Uploading '{file.name}': {current_clause}/{total_paras} clauses...")
                        
                        log_line = f"[{time.strftime('%H:%M:%S')}] [DEBUG] Upserted to Cloud: Clauses {i+1} to {current_clause} from '{file.name}'"
                        logs.append(log_line)
                        log_console.text("\n".join(logs[-12:]))
                        
                    except Exception as err:
                        err_line = f"[{time.strftime('%H:%M:%S')}] ❌ [DEBUG] Error upserting batch {i+1}-{i+len(chunk)}: {str(err)}"
                        logs.append(err_line)
                        log_console.text("\n".join(logs[-12:]))
                        time.sleep(1)
            
            st.success(f"Successfully uploaded {total_indexed} clauses to Pinecone Cloud!")
            time.sleep(1)
            st.rerun()

# ---------------------------------------------------------------------
# TAB 3: CHAT ASSISTANT
# ---------------------------------------------------------------------
with tab_chat:
    st.header("Borrower Contract Chat Assistant")
    
    if "chat_messages" not in st.session_state:
        st.session_state.chat_messages = [
            {"role": "assistant", "content": "Ask me any question regarding borrower-friendly negotiation strategies, precedent loan terms, or project finance risk management."}
        ]
    
    for msg in st.session_state.chat_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            
    if user_query := st.chat_input("Ask a question..."):
        if not nvidia_api_key or not pc_client:
            st.error("Please ensure NVIDIA and Pinecone API Keys are provided.")
        else:
            st.session_state.chat_messages.append({"role": "user", "content": user_query})
            with st.chat_message("user"):
                st.markdown(user_query)
                
            context_str = "NO DIRECT MATCHES FOUND IN PINECONE REPOSITORY."
            try:
                emb_res = pc_client.inference.embed(
                    model=EMBED_MODEL,
                    inputs=[user_query],
                    parameters={"input_type": "query"}
                )
                query_emb = emb_res[0]["values"]
                
                res = pc_index.query(vector=query_emb, top_k=5, include_metadata=True)
                
                context_blocks = []
                if res.get("matches"):
                    for m in res["matches"]:
                        if m["score"] > 0.45:
                            src = m["metadata"].get("source", "Unknown Document")
                            txt = m["metadata"].get("text", "")
                            context_blocks.append(f"From Precedent File [{src}]:\n\"{txt}\"")
                
                if context_blocks:
                    context_str = "\n\n".join(context_blocks)
            except Exception as e:
                print(f"[DEBUG] Chat Error: {e}", flush=True)

            chat_system_prompt = f"""
            You are a Senior Project Finance Legal Advisor representing the BORROWER.
            
            DUAL-LAYER RESPONSE GUIDELINES:
            1. CHECK PINECONE RETRIEVED CONTEXT FIRST:
               - If the query relates to specific baseline terms or definitions stored in your Pinecone Cloud database, explain how to negotiate or maintain those terms to protect the Borrower.
                 
            2. APPLY BORROWER-ADVOCATE LEGAL PRINCIPLES:
               - Provide strategic recommendations on adding cure windows, softening strict covenants, eliminating sole lender discretion, and adding standard borrower carve-outs.

            RETRIEVED PINECONE CLOUD CONTEXT:
            {context_str}
            """
            
            client = OpenAI(
                base_url=NVIDIA_BASE_URL,
                api_key=nvidia_api_key,
                timeout=60.0
            )
            
            try:
                response = client.chat.completions.create(
                    model=NVIDIA_MODEL,
                    messages=[
                        {"role": "system", "content": chat_system_prompt},
                        {"role": "user", "content": user_query}
                    ],
                    temperature=0.2
                )
                answer = response.choices[0].message.content
            except Exception as e:
                answer = f"Error getting response: {str(e)}"
                
            with st.chat_message("assistant"):
                st.markdown(answer)
            st.session_state.chat_messages.append({"role": "assistant", "content": answer})
